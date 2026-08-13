import os
import hashlib
import tempfile
import pytest
from docx import Document
from src.evaluation.evaluate_prospectus import run_evaluation

def test_evaluation_infrastructure_mock(tmp_path):
    mock_input_path = tmp_path / "mock_prospectus.docx"
    mock_report_path = tmp_path / "mock_evaluation_report.md"

    # Create mock prospectus
    doc = Document()
    doc.add_paragraph("Mr. Rajesh Hegde is the Managing Director.")
    doc.add_paragraph("Please contact support at rajesh.hegde@company.com.")
    
    # Save mock file
    doc.save(str(mock_input_path))

    # Record hash to ensure original is unchanged
    with open(str(mock_input_path), "rb") as f:
        original_hash = hashlib.sha256(f.read()).hexdigest()

    # Execute evaluation script
    run_evaluation(str(mock_input_path), str(mock_report_path))

    # Assert original remains byte-for-byte identical
    with open(str(mock_input_path), "rb") as f:
        current_hash = hashlib.sha256(f.read()).hexdigest()
    assert current_hash == original_hash

    # Assert report file exists and is populated
    assert os.path.exists(str(mock_report_path))
    with open(str(mock_report_path), "r", encoding="utf-8") as f:
        report_content = f.read()

    assert "Execution Profile & Runtime Ratios" in report_content
    assert "Document Traversal Summary" in report_content
    assert "Candidate Fusion & Redaction Summary" in report_content
    assert "Redacted PII Category Distribution" in report_content
    assert "PERSON" in report_content
    assert "EMAIL" in report_content
