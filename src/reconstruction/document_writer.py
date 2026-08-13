import os
import shutil
import tempfile
from typing import List, Tuple, Any
from docx import Document
from src.detection.regex_detector import RegexDetector
from src.detection.ner_detector import NERDetector
from src.detection.presidio_detector import PresidioDetector
from src.detection.fusion import resolve_candidates
from src.mapping.span_mapper import reconstruct_paragraph_text, map_span_to_runs
from src.mapping.models import MappedPIISpan
from src.reconstruction.replacer import get_paragraph_runs, apply_replacements
from src.anonymization.entity_store import EntityStore

class RedactionEngine:
    def __init__(self):
        self.regex_detector = RegexDetector()
        self.ner_detector = NERDetector()
        self.presidio_detector = PresidioDetector()

    def get_resolved_entities(self, text: str) -> list:
        """
        Runs detectors and resolves overlaps.
        """
        cands = (
            self.regex_detector.detect(text) +
            self.ner_detector.detect(text) +
            self.presidio_detector.detect(text)
        )
        return resolve_candidates(text, cands)

def redact_document(input_path: str, output_path: str, store: EntityStore) -> list:
    """
    Two-pass DOCX redaction pipeline:
      - Pass 1: Analyzes all text blocks, registers PII candidates globally in EntityStore.
      - Pass 2: Resolves candidates, maps spans, and applies replacements right-to-left.
    Saves atomically using a temp file first.
    
    Returns a redaction manifest: list of dicts with keys
      {paragraph_desc, entity_type, original_text, replacement_text}
    for every entity actually written into the output.
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input file not found at: {input_path}")

    # Load document
    doc = Document(input_path)
    engine = RedactionEngine()

    # Define containers traversal list
    # Every element is a tuple: (paragraph_object, description)
    text_blocks: List[Tuple[Any, str]] = []

    # 1. Body paragraphs
    for idx, p in enumerate(doc.paragraphs):
        text_blocks.append((p, f"body_p_{idx}"))

    # 2. Table cells
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    text_blocks.append((p, f"table_{t_idx}_r_{r_idx}_c_{c_idx}_p_{p_idx}"))

    # 3. Headers and Footers
    for s_idx, section in enumerate(doc.sections):
        if section.header is not None:
            for p_idx, p in enumerate(section.header.paragraphs):
                text_blocks.append((p, f"header_s_{s_idx}_p_{p_idx}"))
        if section.footer is not None:
            for p_idx, p in enumerate(section.footer.paragraphs):
                text_blocks.append((p, f"footer_s_{s_idx}_p_{p_idx}"))

    # ==========================================================================
    # PASS 1: Candidate Detection & Entity Registration
    # ==========================================================================
    all_resolved_pii: List[Tuple[Any, list, list, str]] = []  # (paragraph, runs, entities, desc)
    seen_paragraphs = set()
    
    for paragraph, desc in text_blocks:
        p_key = paragraph._p
        if p_key in seen_paragraphs:
            continue
        seen_paragraphs.add(p_key)
        
        runs = get_paragraph_runs(paragraph)
        if not runs:
            continue
        text, offsets = reconstruct_paragraph_text(runs)
        if not text.strip():
            continue
            
        resolved = engine.get_resolved_entities(text)
        # Filter out generic DATE (non-PII) to prevent redacting generic document dates
        resolved = [e for e in resolved if e.entity_type != "DATE"]
        
        # Candidate Validation Layer
        from src.detection.validation import CandidateValidator
        validator = CandidateValidator()
        validated = validator.validate_candidates(text, resolved)
        
        # Only keep candidates approved by the validation layer
        keep_pii = [e for e in validated if e.metadata.get("validation_decision") == "KEEP"]
        
        if keep_pii:
            # Register in EntityStore
            store.register_candidates(keep_pii)
            all_resolved_pii.append((paragraph, runs, keep_pii, desc))

    # Generate synthetic replacements for registered entities
    store.generate_all_replacements()

    # Compile propagation list
    import re
    from src.detection.models import PIIEntity
    propagation_entities = []
    
    corporate_suffixes = {"limited", "ltd", "ltd.", "private", "pvt", "llp", "corporation", "inc", "inc.", "plc", "llc", "group"}
    
    for canonical_id, orig_text in store.id_to_original.items():
        entity_type = canonical_id.split("_")[0]
        words = orig_text.split()
        num_words = len(words)
        
        should_propagate = False
        
        if entity_type == "PERSON":
            if num_words >= 2:
                should_propagate = True
        elif entity_type == "COMPANY":
            has_suffix = any(w.lower() in corporate_suffixes for w in words)
            if num_words >= 2 or has_suffix:
                should_propagate = True
        elif entity_type in {"EMAIL", "PHONE", "CREDIT_CARD", "SSN", "IP_ADDRESS"}:
            should_propagate = True
            
        if should_propagate:
            # Use word boundaries if it's alphanumeric at the edges.
            esc_text = re.escape(orig_text)
            pattern_str = esc_text
            if orig_text[0].isalnum():
                pattern_str = r"\b" + pattern_str
            if orig_text[-1].isalnum():
                pattern_str = pattern_str + r"\b"
                
            propagation_entities.append({
                "type": entity_type,
                "text": orig_text,
                "pattern": re.compile(pattern_str, re.IGNORECASE)
            })

    # ==========================================================================
    # PASS 2: Span Mapping & Replacement Application
    # ==========================================================================
    redaction_manifest = []
    local_entities_map = {p_obj._p: ents for p_obj, r, ents, d in all_resolved_pii}
    
    seen_paragraphs_pass2 = set()
    for paragraph, desc in text_blocks:
        p_key = paragraph._p
        if p_key in seen_paragraphs_pass2:
            continue
        seen_paragraphs_pass2.add(p_key)
        
        runs = get_paragraph_runs(paragraph)
        if not runs:
            continue
        text, offsets = reconstruct_paragraph_text(runs)
        if not text.strip():
            continue
            
        local_entities = local_entities_map.get(p_key, [])
        propagated_cands = []
        
        for prop in propagation_entities:
            for match in prop["pattern"].finditer(text):
                ent = PIIEntity(
                    entity_type=prop["type"],
                    text=match.group(),
                    start=match.start(),
                    end=match.end(),
                    confidence=1.0,
                    source="propagation"
                )
                propagated_cands.append(ent)
                
        if not local_entities and not propagated_cands:
            continue
            
        combined_cands = local_entities + propagated_cands
        from src.detection.fusion import resolve_candidates
        final_resolved = resolve_candidates(text, combined_cands)
        
        mapped_spans: List[MappedPIISpan] = []
        
        for entity in final_resolved:
            # Map logical char offsets to runs
            run_spans = map_span_to_runs(entity.start, entity.end, entity, offsets, runs)
            mapped_spans.append(MappedPIISpan(entity, run_spans))
            
            # Record in manifest
            replacement = store.get_replacement(entity.entity_type, entity.text)
            redaction_manifest.append({
                "paragraph_desc": desc,
                "entity_type": entity.entity_type,
                "original_text": entity.text,
                "replacement_text": replacement,
            })
            
        # Apply replacement
        apply_replacements(paragraph, mapped_spans, runs, store)

    # ==========================================================================
    # ATOMIC SAVE finalization
    # ==========================================================================
    temp_dir = tempfile.gettempdir()
    temp_file_path = os.path.join(temp_dir, f"redacted_temp_{os.urandom(8).hex()}.docx")
    
    try:
        doc.save(temp_file_path)
        
        # Verify temporary file compiles / is readable
        verify_doc = Document(temp_file_path)
        
        # Atomically write to final destination
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        shutil.move(temp_file_path, output_path)
    except Exception as e:
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
        raise RuntimeError(f"Failed to safely reconstruct/save redacted document: {e}")
    
    return redaction_manifest

