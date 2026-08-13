import os
import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.detection.models import PIIEntity
from src.anonymization.generator import SyntheticGenerator
from src.anonymization.entity_store import EntityStore
from src.reconstruction.document_writer import redact_document

def test_adversarial_split_runs_and_formatting(tmp_path):
    input_doc_path = tmp_path / "split_runs.docx"
    output_doc_path = tmp_path / "split_runs_redacted.docx"
    
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("The Whole-time Director ")
    p.add_run("Raj").bold = True
    p.add_run("esh").italic = True
    p.add_run(" Hegde")
    p.add_run(" was present.")
    
    doc.save(str(input_doc_path))
    
    generator = SyntheticGenerator(locale="en_IN")
    store = EntityStore(generator)
    redact_document(str(input_doc_path), str(output_doc_path), store)
    
    red_doc = Document(str(output_doc_path))
    red_text = red_doc.paragraphs[0].text
    
    assert "Rajesh Hegde" not in red_text
    fake_name = store.get_replacement("PERSON", "Rajesh Hegde")
    assert fake_name in red_text
    
    # Assert bold/italic formatting runs survived replacement
    runs = red_doc.paragraphs[0].runs
    assert any(r.bold for r in runs)
    assert any(r.italic for r in runs)

def test_adversarial_tabs_and_nbsp(tmp_path):
    input_doc_path = tmp_path / "tabs_nbsp.docx"
    output_doc_path = tmp_path / "tabs_nbsp_redacted.docx"
    
    doc = Document()
    doc.add_paragraph("Report issued by ICICI\tSecurities\xa0Limited in Mumbai.")
    doc.save(str(input_doc_path))
    
    generator = SyntheticGenerator(locale="en_IN")
    store = EntityStore(generator)
    redact_document(str(input_doc_path), str(output_doc_path), store)
    
    red_doc = Document(str(output_doc_path))
    red_text = red_doc.paragraphs[0].text
    
    # Raw tab/nbsp get normalized to spaces during reconstruction
    # Assert we successfully redacted the normalized name
    assert "ICICI Securities Limited" not in red_text
    fake_comp = store.get_replacement("COMPANY", "ICICI Securities Limited")
    assert fake_comp in red_text

def test_adversarial_punctuation_and_parentheses(tmp_path):
    input_doc_path = tmp_path / "punctuation.docx"
    output_doc_path = tmp_path / "punctuation_redacted.docx"
    
    doc = Document()
    doc.add_paragraph("Contact: [Rajesh Hegde], (Chairman).")
    doc.save(str(input_doc_path))
    
    generator = SyntheticGenerator(locale="en_IN")
    store = EntityStore(generator)
    redact_document(str(input_doc_path), str(output_doc_path), store)
    
    red_doc = Document(str(output_doc_path))
    red_text = red_doc.paragraphs[0].text
    
    assert "Rajesh Hegde" not in red_text
    fake_name = store.get_replacement("PERSON", "Rajesh Hegde")
    assert f"[{fake_name}]" in red_text

def test_adversarial_repeated_aliases(tmp_path):
    input_doc_path = tmp_path / "aliases.docx"
    output_doc_path = tmp_path / "aliases_redacted.docx"
    
    doc = Document()
    doc.add_paragraph("Director Rajesh Hegde attended.")
    doc.add_paragraph("Please contact Rajesh Hegde immediately.")
    doc.save(str(input_doc_path))
    
    generator = SyntheticGenerator(locale="en_IN")
    store = EntityStore(generator)
    redact_document(str(input_doc_path), str(output_doc_path), store)
    
    red_doc = Document(str(output_doc_path))
    text1 = red_doc.paragraphs[0].text
    text2 = red_doc.paragraphs[1].text
    
    assert "Rajesh Hegde" not in text1
    assert "Rajesh Hegde" not in text2
    
    fake1 = store.get_replacement("PERSON", "Rajesh Hegde")
    fake2 = store.get_replacement("PERSON", "Rajesh Hegde")
    
    # Assert synthetic names are identical across paragraphs
    assert fake1 == fake2
    assert fake1 in text1
    assert fake1 in text2

def test_adversarial_nested_and_overlapping(tmp_path):
    input_doc_path = tmp_path / "nested.docx"
    output_doc_path = tmp_path / "nested_redacted.docx"
    
    doc = Document()
    doc.add_paragraph("Registered at Registrar of Companies, Maharashtra.")
    doc.save(str(input_doc_path))
    
    generator = SyntheticGenerator(locale="en_IN")
    store = EntityStore(generator)
    redact_document(str(input_doc_path), str(output_doc_path), store)
    
    red_doc = Document(str(output_doc_path))
    red_text = red_doc.paragraphs[0].text
    
    assert "Registrar of Companies, Maharashtra" not in red_text
    fake_comp = store.get_replacement("COMPANY", "Registrar of Companies, Maharashtra")
    assert fake_comp in red_text
