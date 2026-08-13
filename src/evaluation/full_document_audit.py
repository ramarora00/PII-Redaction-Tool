import os
import re
import json
from collections import Counter, defaultdict
from typing import List, Dict, Any, Tuple
from docx import Document

from src.detection.regex_detector import RegexDetector
from src.detection.ner_detector import NERDetector
from src.detection.presidio_detector import PresidioDetector
from src.detection.fusion import resolve_candidates
from src.detection.validation import CandidateValidator
from src.mapping.span_mapper import reconstruct_paragraph_text
from src.reconstruction.replacer import get_paragraph_runs

def run_full_document_audit(doc_path: str, output_report_path: str) -> None:
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Prospectus input file not found at: {doc_path}")

    print(f"Starting full-document PII audit over: {doc_path}")
    doc = Document(doc_path)

    regex_detector = RegexDetector()
    ner_detector = NERDetector()
    presidio_detector = PresidioDetector()
    candidate_validator = CandidateValidator()

    # Track overall block counts
    block_counts = Counter()
    
    # 1. Validation stats
    keep_counts = Counter()
    reject_counts = Counter()
    rejection_reasons = Counter()
    
    # 2. Detector disagreement stats
    disagreement_counts = Counter()
    
    # 3. Discovered heuristic FNs
    discovered_fns = []
    
    # 4. FP patterns clustering
    fp_clusters = defaultdict(list)

    # Reconstruct prospectus blocks
    text_blocks = []
    for idx, p in enumerate(doc.paragraphs):
        text_blocks.append((p, "body", f"body / paragraph={idx}"))
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    text_blocks.append((p, "table", f"table / table={t_idx} / row={r_idx} / cell={c_idx} / paragraph={p_idx}"))
    for s_idx, section in enumerate(doc.sections):
        if section.header is not None:
            for p_idx, p in enumerate(section.header.paragraphs):
                text_blocks.append((p, "header", f"header / section={s_idx} / paragraph={p_idx}"))
        if section.footer is not None:
            for p_idx, p in enumerate(section.footer.paragraphs):
                text_blocks.append((p, "footer", f"footer / section={s_idx} / paragraph={p_idx}"))

    total_blocks = len(text_blocks)
    print(f"Scanned prospectus structure: {total_blocks} blocks found.")

    # Missed PII patterns (Heuristic FNs search)
    missed_person_pattern = re.compile(r"\b(?:Mr\.|Ms\.|Mrs\.|Shri|Shree|Dr\.)\s+([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b")
    missed_company_pattern = re.compile(r"\b([A-Z][a-zA-Z0-9&]*(?:\s+[A-Z][a-zA-Z0-9&]*){1,4}\s+(?:Limited|Private Limited|LLP|Corporation|Pvt\.\s+Ltd\.|Ltd\.))\b")
    missed_email_pattern = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b")
    missed_phone_pattern = re.compile(r"\b(?:\+91|91)?[6789]\d{9}\b")

    for idx, (p, label, path) in enumerate(text_blocks):
        block_counts[label] += 1
        runs = get_paragraph_runs(p)
        if not runs:
            continue
        text, offsets = reconstruct_paragraph_text(runs)
        if not text.strip():
            continue

        if idx > 0 and idx % 1000 == 0:
            print(f"Audited {idx}/{total_blocks} blocks...")

        # Run detection
        regex_cands = regex_detector.detect(text)
        ner_cands = ner_detector.detect(text)
        presidio_cands = presidio_detector.detect(text)
        
        all_raw_cands = regex_cands + ner_cands + presidio_cands
        resolved = resolve_candidates(text, all_raw_cands)
        resolved_pii = [e for e in resolved if e.entity_type != "DATE"]
        validated = candidate_validator.validate_candidates(text, resolved_pii)

        for cand in validated:
            decision = cand.metadata.get("validation_decision", "KEEP")
            reason = cand.metadata.get("validation_reason", "GENUINE_PII")
            
            # Trace provenance to analyze detector disagreement
            # Find which detectors originally generated this candidate's span
            sources = set()
            for r_cand in all_raw_cands:
                # If spans overlap or match closely, compile sources
                if max(r_cand.start, cand.start) < min(r_cand.end, cand.end):
                    sources.add(r_cand.source)
            
            sources_sorted = sorted(list(sources))
            prov_key = " + ".join(sources_sorted) if sources_sorted else "unknown"
            
            if decision == "KEEP":
                keep_counts[cand.entity_type] += 1
                disagreement_counts[prov_key] += 1
            else:
                reject_counts[cand.entity_type] += 1
                rejection_reasons[reason] += 1
                
                # FP Clustering: Save rejected entities by reason for pattern categorization
                fp_clusters[reason].append((cand.text, cand.entity_type, path))

        # Heuristic FN Discovery (unmatched PII search)
        def is_covered(start: int, end: int) -> bool:
            return any(max(start, c.start) < min(end, c.end) for c in validated)

        # 1. Missed PERSON names
        for m in missed_person_pattern.finditer(text):
            if not is_covered(m.start(), m.end()):
                discovered_fns.append({
                    "text": m.group(0),
                    "type": "PERSON",
                    "path": path,
                    "context": text[max(0, m.start()-30):min(len(text), m.end()+30)]
                })

        # 2. Missed COMPANY names
        for m in missed_company_pattern.finditer(text):
            if not is_covered(m.start(), m.end()):
                discovered_fns.append({
                    "text": m.group(0),
                    "type": "COMPANY",
                    "path": path,
                    "context": text[max(0, m.start()-30):min(len(text), m.end()+30)]
                })

        # 3. Missed EMAILs
        for m in missed_email_pattern.finditer(text):
            if not is_covered(m.start(), m.end()):
                discovered_fns.append({
                    "text": m.group(0),
                    "type": "EMAIL",
                    "path": path,
                    "context": text[max(0, m.start()-30):min(len(text), m.end()+30)]
                })

        # 4. Missed PHONEs
        for m in missed_phone_pattern.finditer(text):
            if not is_covered(m.start(), m.end()):
                discovered_fns.append({
                    "text": m.group(0),
                    "type": "PHONE",
                    "path": path,
                    "context": text[max(0, m.start()-30):min(len(text), m.end()+30)]
                })

    # Prepare clustering and disagreement tables
    disagreement_rows = []
    for combo, count in disagreement_counts.most_common():
        disagreement_rows.append(f"| {combo} | {count} |")

    reason_rows = []
    for reason, count in rejection_reasons.most_common():
        reason_rows.append(f"| `{reason}` | {count} |")

    # Group FP patterns into clusters
    clustered_patterns = []
    for reason, items in fp_clusters.items():
        sample_items = Counter([it[0] for it in items]).most_common(10)
        samples_str = ", ".join([f'"{it[0]}" ({it[1]})' for it in sample_items])
        clustered_patterns.append(f"| `{reason}` | {len(items)} | {samples_str} |")

    # Generate full report
    os.makedirs(os.path.dirname(os.path.abspath(output_report_path)), exist_ok=True)
    with open(output_report_path, "w", encoding="utf-8") as f:
        f.write(f"""# Full-Document PII Audit Report

This report presents diagnostic profiling statistics and error clusters collected across the entire `prospectus.docx` (scanned {total_blocks} structural blocks).

---

## 1. Document Structure & Coverage Summary
- **Scanned Blocks by Container**:
  - Body Paragraphs: {block_counts['body']}
  - Table Cell Paragraphs: {block_counts['table']}
  - Section Headers: {block_counts['header']}
  - Section Footers: {block_counts['footer']}

- **Validated (KEPT) PII Entities by Type**:
""" + "\n".join([f"  - `{k}`: {v}" for k, v in keep_counts.items()]) + f"""

- **Validated (REJECTED) PII Entities by Type**:
""" + "\n".join([f"  - `{k}`: {v}" for k, v in reject_counts.items()]) + f"""

---

## 2. Cross-Detector Disagreement Analysis
This table shows kept candidate counts matching each unique combination of detector source provenance.

| Detector Combination | Count |
| :--- | :---: |
""" + "\n".join(disagreement_rows) + f"""

---

## 3. False Positive Clustering (Validator Rejections)
Candidates rejected by the validation layer clustered by rejection pattern:

| Rejection Reason | Total Rejections | Sample Cluster Items (Count) |
| :--- | :---: | :--- |
""" + "\n".join(clustered_patterns) + f"""

---

## 4. Discovered Heuristic FNs (Missed PII Search)
The following potential PII elements matched name/structure heuristics but were not detected by the pipeline.

Total Discovered Heuristic FNs: {len(discovered_fns)}

| Entity Type | Discovered Text | Location Block | Sentence Context |
| :--- | :--- | :--- | :--- |
""" + "\n".join([f"| `{fn['type']}` | `\"{fn['text']}\"` | `{fn['path']}` | \"...{fn['context'].strip().replace(chr(10), ' ')}...\" |" for fn in discovered_fns[:25]]) + """
""")

    print(f"Full document audit report successfully written to: {output_report_path}")

if __name__ == "__main__":
    run_full_document_audit("data/input/prospectus.docx", "evaluation/full_document_audit_report.md")
