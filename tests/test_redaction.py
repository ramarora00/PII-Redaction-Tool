import os
import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from src.detection.models import PIIEntity
from src.anonymization.generator import SyntheticGenerator
from src.anonymization.entity_store import EntityStore
from src.reconstruction.replacer import get_paragraph_runs, apply_replacements
from src.reconstruction.document_writer import redact_document

class MockRun:
    def __init__(self, text: str, bold: bool = False):
        self.text = text
        self.bold = bold

def test_paragraph_runs_hyperlink_traversal():
    doc = Document()
    p = doc.add_paragraph("Go to ")
    
    # Manually append hyperlink tag to paragraph XML element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), "rId1")
    
    r_node = OxmlElement("w:r")
    r_text = OxmlElement("w:t")
    r_text.text = "www.google.com"
    r_node.append(r_text)
    hyperlink.append(r_node)
    
    p._p.append(hyperlink)
    
    # We should detect 2 runs total (one body run, one hyperlink run)
    runs = get_paragraph_runs(p)
    assert len(runs) == 2
    assert runs[0].text == "Go to "
    assert runs[1].text == "www.google.com"

def test_inline_replacement_single_run():
    runs = [MockRun("Chairman Rajesh Hegde was present.", bold=True)]
    generator = SyntheticGenerator(locale="en_IN")
    store = EntityStore(generator)
    
    entity = PIIEntity("PERSON", "Rajesh Hegde", 9, 21, 0.90, "presidio")
    store.register_candidates([entity])
    store.generate_all_replacements()
    
    fake_name = store.get_replacement("PERSON", "Rajesh Hegde")
    
    from src.mapping.span_mapper import reconstruct_paragraph_text, map_span_to_runs
    from src.mapping.models import MappedPIISpan
    
    text, offsets = reconstruct_paragraph_text(runs)
    run_spans = map_span_to_runs(9, 21, entity, offsets, runs)
    mapped_spans = [MappedPIISpan(entity, run_spans)]
    
    apply_replacements(None, mapped_spans, runs, store)
    
    assert fake_name in runs[0].text
    assert "Chairman " in runs[0].text
    assert " was present." in runs[0].text
    assert runs[0].bold is True  # preserved formatting

def test_inline_replacement_cross_run():
    runs = [
        MockRun("Chairman "),
        MockRun("Raj"),
        MockRun("esh Hegde"),
        MockRun(" was present.")
    ]
    generator = SyntheticGenerator(locale="en_IN")
    store = EntityStore(generator)
    
    entity = PIIEntity("PERSON", "Rajesh Hegde", 9, 21, 0.90, "presidio")
    store.register_candidates([entity])
    store.generate_all_replacements()
    
    fake_name = store.get_replacement("PERSON", "Rajesh Hegde")
    
    from src.mapping.span_mapper import reconstruct_paragraph_text, map_span_to_runs
    from src.mapping.models import MappedPIISpan
    
    text, offsets = reconstruct_paragraph_text(runs)
    run_spans = map_span_to_runs(9, 21, entity, offsets, runs)
    mapped_spans = [MappedPIISpan(entity, run_spans)]
    
    apply_replacements(None, mapped_spans, runs, store)
    
    # First matched run receives the replacement name
    assert runs[1].text == fake_name
    # Second matched run matched segment is cleared
    assert runs[2].text == ""
    # Standard text remains unchanged
    assert runs[0].text == "Chairman "
    assert runs[3].text == " was present."

def test_e2e_document_redaction(tmp_path):
    input_doc_path = tmp_path / "prospectus_input.docx"
    output_doc_path = tmp_path / "prospectus_redacted.docx"
    
    # Create input document
    doc = Document()
    
    # 1. Standard body text
    p1 = doc.add_paragraph("Mr. Rajesh Hegde is the Chairman of the Board.")
    
    # 2. Table with columns containing PII
    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Director"
    hdr_cells[1].text = "Email"
    
    row_cells = table.rows[1].cells
    row_cells[0].text = "Rajesh Hegde"
    row_cells[1].text = "rajesh.hegde@company.com"
    
    # 3. Header/Footer
    section = doc.sections[0]
    section.header.paragraphs[0].text = "Confidential - Date of Birth: 15-05-1990"
    section.footer.paragraphs[0].text = "Page 1 - Date of Issue: 2026/08/13"  # generic date should NOT be redacted
    
    # 4. Hyperlink element XML insertion
    p_link = doc.add_paragraph("Send credentials to ")
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), "rId99")
    
    r_node = OxmlElement("w:r")
    r_text = OxmlElement("w:t")
    r_text.text = "rajesh.hegde@company.com"
    r_node.append(r_text)
    hyperlink.append(r_node)
    p_link._p.append(hyperlink)
    
    # Save original document
    doc.save(str(input_doc_path))
    
    # Get hash of original document to verify immutable status
    import hashlib
    with open(str(input_doc_path), "rb") as f:
        original_hash = hashlib.sha256(f.read()).hexdigest()
        
    # Execute Redaction
    generator = SyntheticGenerator()
    store = EntityStore(generator)
    
    redact_document(str(input_doc_path), str(output_doc_path), store)
    
    # Assert original file is byte-for-byte unchanged
    with open(str(input_doc_path), "rb") as f:
        current_hash = hashlib.sha256(f.read()).hexdigest()
    assert current_hash == original_hash
    
    # Verify redacted document
    redacted_doc = Document(str(output_doc_path))
    
    # 1. Body name is replaced
    p1_red = redacted_doc.paragraphs[0]
    assert "Rajesh Hegde" not in p1_red.text
    fake_name = store.get_replacement("PERSON", "Rajesh Hegde")
    assert fake_name in p1_red.text
    
    # 2. Table cells are redacted
    table_red = redacted_doc.tables[0]
    assert table_red.rows[1].cells[0].text == fake_name
    fake_email = store.get_replacement("EMAIL", "rajesh.hegde@company.com")
    assert table_red.rows[1].cells[1].text == fake_email
    
    # 3. Header PII is replaced
    header_p = redacted_doc.sections[0].header.paragraphs[0]
    assert "15-05-1990" not in header_p.text
    fake_dob = store.get_replacement("DATE_OF_BIRTH", "15-05-1990")
    assert fake_dob in header_p.text
    
    # 4. Footer generic date is NOT redacted
    footer_p = redacted_doc.sections[0].footer.paragraphs[0]
    assert "2026/08/13" in footer_p.text  # remains unmodified
    
    # 5. Hyperlink run is mapped and replaced
    # Since there are 2 body paragraphs, index of p_link is 1 (paragraphs[1])
    p_link_red = redacted_doc.paragraphs[1]
    assert "rajesh.hegde@company.com" not in p_link_red.text
    assert fake_email in p_link_red.text
