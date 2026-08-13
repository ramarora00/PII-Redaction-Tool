import os
import pytest
from docx import Document
from src.doc_inspector import DocInspector

@pytest.fixture
def dummy_docx(tmp_path):
    doc_path = tmp_path / "test_doc.docx"
    doc = Document()
    
    # 1. Add simple paragraphs (some with multiple runs)
    p1 = doc.add_paragraph()
    p1.add_run("This is the first paragraph.")
    
    # 2. Add paragraph with split word (fragmentation)
    p2 = doc.add_paragraph()
    p2.add_run("The Whole-time Dire")
    p2.add_run("ctor ")  # Split between runs without space
    p2.add_run("is Mr. John.")

    
    # 3. Add table
    table = doc.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Header 1'
    hdr_cells[1].text = 'Header 2'
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Director info'
    row_cells[1].text = 'email: director@company.com'
    
    doc.save(str(doc_path))
    return str(doc_path)

def test_basic_statistics(dummy_docx):
    inspector = DocInspector(dummy_docx)
    stats = inspector.get_basic_stats()
    
    assert stats["total_primary_paragraphs"] == 2
    assert stats["total_tables"] == 1
    assert stats["total_rows"] == 2
    assert stats["total_cells"] == 4
    # The cell assignment table.rows[x].cells[y].text sets paragraphs implicitly
    assert stats["cell_paragraphs_count"] == 4

def test_run_fragmentation(dummy_docx):
    inspector = DocInspector(dummy_docx)
    fragmented = inspector.analyze_run_fragmentation()
    
    assert len(fragmented) > 0
    first_frag = fragmented[0]
    assert "Dire" in first_frag["run_text_1"]
    assert "ctor" in first_frag["run_text_2"]


def test_find_sample_pii_locations(dummy_docx):
    inspector = DocInspector(dummy_docx)
    locations = inspector.find_sample_pii_locations()
    
    # We added "Director" and "email: director@company.com"
    # These contain keywords "director" and "email"
    assert len(locations) > 0
    loc_names = [l["location"] for l in locations]
    # At least one should be in paragraph or table cell
    assert any("paragraph" in name or "table" in name for name in loc_names)
