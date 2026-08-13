import os
from typing import Dict, List, Any, Tuple
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

class DocInspector:
    def __init__(self, docx_path: str):
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found at: {docx_path}")
        self.docx_path = docx_path
        self.doc = Document(docx_path)

    def get_basic_stats(self) -> Dict[str, Any]:
        """
        Gathers basic statistics of the document.
        """
        # Count primary paragraphs (non-table, non-header/footer)
        primary_paragraphs = self.doc.paragraphs
        total_primary_paragraphs = len(primary_paragraphs)
        
        # Count runs in primary paragraphs
        primary_runs_count = sum(len(p.runs) for p in primary_paragraphs)

        # Count tables, rows, cells, and paragraphs in cells
        tables = self.doc.tables
        total_tables = len(tables)
        total_rows = 0
        total_cells = 0
        cell_paragraphs_count = 0
        cell_runs_count = 0

        for table in tables:
            for row in table.rows:
                total_rows += 1
                for cell in row.cells:
                    total_cells += 1
                    cell_paragraphs_count += len(cell.paragraphs)
                    for cp in cell.paragraphs:
                        cell_runs_count += len(cp.runs)

        # Count headers and footers (across all sections)
        header_paragraphs_count = 0
        header_runs_count = 0
        footer_paragraphs_count = 0
        footer_runs_count = 0

        for section in self.doc.sections:
            if section.header is not None:
                header_paragraphs_count += len(section.header.paragraphs)
                header_runs_count += sum(len(p.runs) for p in section.header.paragraphs)
            if section.footer is not None:
                footer_paragraphs_count += len(section.footer.paragraphs)
                footer_runs_count += sum(len(p.runs) for p in section.footer.paragraphs)

        return {
            "total_primary_paragraphs": total_primary_paragraphs,
            "primary_runs_count": primary_runs_count,
            "total_tables": total_tables,
            "total_rows": total_rows,
            "total_cells": total_cells,
            "cell_paragraphs_count": cell_paragraphs_count,
            "cell_runs_count": cell_runs_count,
            "header_paragraphs_count": header_paragraphs_count,
            "header_runs_count": header_runs_count,
            "footer_paragraphs_count": footer_paragraphs_count,
            "footer_runs_count": footer_runs_count,
        }

    def analyze_run_fragmentation(self) -> List[Dict[str, Any]]:
        """
        Finds instances in paragraphs where text is split across multiple runs.
        We check if contiguous words/patterns are divided.
        """
        fragmented_examples = []
        # Examine first 100 paragraphs for fragmentation examples to avoid huge logs
        for idx, p in enumerate(self.doc.paragraphs[:200]):
            if len(p.runs) <= 1:
                continue
            
            # Simple check: if a run doesn't end with a space or punctuation, or if adjacent runs are very short
            # (e.g. single character runs, or runs containing part of a word).
            # Let's inspect the actual runs.
            run_texts = [r.text for r in p.runs]
            full_text = p.text
            
            # Look for split words: e.g., run_texts has "Direct", "or" but the word is "Director"
            # We can find where a run ends mid-word.
            for i in range(len(p.runs) - 1):
                r1 = p.runs[i].text
                r2 = p.runs[i+1].text
                if not r1 or not r2:
                    continue
                # If r1 ends with an alphanumeric character and r2 starts with one, the word is split across runs!
                if r1[-1].isalnum() and r2[0].isalnum():
                    # Find the word context
                    word_part1 = r1.split()[-1] if r1.split() else r1
                    word_part2 = r2.split()[0] if r2.split() else r2
                    reconstructed_word = word_part1 + word_part2
                    fragmented_examples.append({
                        "paragraph_index": idx,
                        "run_index": i,
                        "run_text_1": r1,
                        "run_text_2": r2,
                        "reconstructed_word": reconstructed_word,
                        "context": f"...{r1[-15:]} | {r2[:15]}..."
                    })
                    if len(fragmented_examples) >= 10:
                        break
            if len(fragmented_examples) >= 10:
                break
                
        return fragmented_examples

    def find_sample_pii_locations(self) -> List[Dict[str, Any]]:
        """
        Scans paragraphs and tables for patterns resembling PII (like email, phone, or potential names)
        to identify where they occur (e.g. headers, footers, paragraphs, or table cells).
        Does NOT extract or dump the actual PII content itself.
        """
        occurrences = []
        
        # Simple keywords indicating PII containers in financial documents
        keywords = ["director", "email", "phone", "telephone", "address", "registered office", "promoter", "pan"]

        # Check primary paragraphs
        for idx, p in enumerate(self.doc.paragraphs):
            text_lower = p.text.lower()
            for kw in keywords:
                if kw in text_lower:
                    occurrences.append({
                        "location": f"paragraph[{idx}]",
                        "context_hint": f"Contains keyword '{kw}'",
                        "character_length": len(p.text)
                    })
                    break
            if len(occurrences) >= 10:
                break

        # Check tables
        for t_idx, table in enumerate(self.doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, p in enumerate(cell.paragraphs):
                        text_lower = p.text.lower()
                        for kw in keywords:
                            if kw in text_lower:
                                occurrences.append({
                                    "location": f"table[{t_idx}].row[{r_idx}].cell[{c_idx}].paragraph[{p_idx}]",
                                    "context_hint": f"Table cell contains keyword '{kw}'",
                                    "character_length": len(p.text)
                                })
                                break
                    if len(occurrences) >= 20:
                        break
                if len(occurrences) >= 20:
                    break
            if len(occurrences) >= 20:
                break

        return occurrences
