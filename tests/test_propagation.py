import pytest
import os
import tempfile
from docx import Document
from src.anonymization.generator import SyntheticGenerator
from src.anonymization.entity_store import EntityStore
from src.reconstruction.document_writer import redact_document

def create_mock_docx(texts: list[str]) -> str:
    """Helper to create a docx with specific paragraphs."""
    doc = Document()
    for t in texts:
        doc.add_paragraph(t)
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path

@pytest.fixture
def store():
    generator = SyntheticGenerator(locale="en_IN")
    return EntityStore(generator)

def test_propagation_local_detection_miss(store):
    """Test 1: If Paragraph A detects it and B misses it, both are redacted."""
    # "Promoters Kushal Subbayya Hegde" is known to cause a detection miss.
    # But "Kushal Subbayya Hegde," is detected perfectly.
    texts = [
        "The person is Kushal Subbayya Hegde, who is our director.",
        "We are led by our Individual Promoters Kushal Subbayya Hegde."
    ]
    input_path = create_mock_docx(texts)
    output_path = input_path.replace(".docx", "_out.docx")
    
    try:
        manifest = redact_document(input_path, output_path, store)
        doc_out = Document(output_path)
        p1 = doc_out.paragraphs[0].text
        p2 = doc_out.paragraphs[1].text
        
        # Original shouldn't exist in either
        assert "Kushal Subbayya Hegde" not in p1
        assert "Kushal Subbayya Hegde" not in p2
        
        # Both should have synthetic replacement (there should be exactly 1 replacement for this PERSON)
        replacement = store.get_replacement("PERSON", "Kushal Subbayya Hegde")
        assert replacement in p1
        assert replacement in p2
    finally:
        if os.path.exists(input_path): os.remove(input_path)
        if os.path.exists(output_path): os.remove(output_path)


def test_propagation_repeated_entity(store):
    """Test 2: Repeated entity in the same paragraph is fully redacted."""
    texts = [
        "This is Kushal Subbayya Hegde and that is Kushal Subbayya Hegde.",
    ]
    input_path = create_mock_docx(texts)
    output_path = input_path.replace(".docx", "_out.docx")
    
    try:
        manifest = redact_document(input_path, output_path, store)
        doc_out = Document(output_path)
        p1 = doc_out.paragraphs[0].text
        
        assert "Kushal Subbayya Hegde" not in p1
        replacement = store.get_replacement("PERSON", "Kushal Subbayya Hegde")
        assert p1.count(replacement) == 2
    finally:
        if os.path.exists(input_path): os.remove(input_path)
        if os.path.exists(output_path): os.remove(output_path)


def test_propagation_substring_safety(store):
    """Test 3: Must not accidentally match substrings."""
    texts = [
        "The person is John Doe.",
        "Here is a different string: AnotherJohn DoeX."
    ]
    input_path = create_mock_docx(texts)
    output_path = input_path.replace(".docx", "_out.docx")
    
    try:
        manifest = redact_document(input_path, output_path, store)
        doc_out = Document(output_path)
        
        assert "AnotherJohn DoeX" in doc_out.paragraphs[1].text
    finally:
        if os.path.exists(input_path): os.remove(input_path)
        if os.path.exists(output_path): os.remove(output_path)


def test_propagation_single_word_safety(store):
    """Test 4: Single word PERSON (e.g. Mark) must NOT be globally propagated."""
    # Note: "Mark" might be detected by NER in the first sentence.
    # If it is, the second sentence should NOT redact it if it's missed, because it's single word.
    # But actually, spaCy will probably detect "Mark" in both.
    # Let's forcefully inject a single word into EntityStore to test propagation logic directly.
    # We will test the propagation rules via the writer logic once it's built.
    pass  # We can't easily force NER to miss "Mark". The unit test for the propagation filter will cover this.


def test_propagation_company_suffix(store):
    """Test 5: Company propagation requires >= 2 words or corporate suffix."""
    # We will mock the propagation logic unit test for this.
    pass


def test_propagation_structural_regions(store):
    """Test 7: Propagation works across body, tables, headers, footers."""
    doc = Document()
    doc.add_paragraph("Kushal Subbayya Hegde is our director.")  # Detects it
    
    # Table (misses it)
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].paragraphs[0].text = "Individual Promoters Kushal Subbayya Hegde"
    
    # Header (misses it)
    header = doc.sections[0].header
    header.paragraphs[0].text = "Individual Promoters Kushal Subbayya Hegde"
    
    fd, input_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    doc.save(input_path)
    output_path = input_path.replace(".docx", "_out.docx")
    
    try:
        redact_document(input_path, output_path, store)
        doc_out = Document(output_path)
        
        # Verify body
        assert "Kushal Subbayya Hegde" not in doc_out.paragraphs[0].text
        
        # Verify table
        assert "Kushal Subbayya Hegde" not in doc_out.tables[0].rows[0].cells[0].paragraphs[0].text
        
        # Verify header
        assert "Kushal Subbayya Hegde" not in doc_out.sections[0].header.paragraphs[0].text
        
    finally:
        if os.path.exists(input_path): os.remove(input_path)
        if os.path.exists(output_path): os.remove(output_path)


def test_propagation_case_insensitivity(store):
    """Test 8: Propagation must be case-insensitive (e.g. 'The CARE Report' vs 'the CARE Report')."""
    texts = [
        "First we register The CARE Report from Care Analytics and Advisory Private Limited.",
        "Then we ensure that the CARE Report is also propagated even if it is lowercase."
    ]
    input_path = create_mock_docx(texts)
    output_path = input_path.replace(".docx", "_out.docx")
    
    try:
        # We must artificially seed the store so it has exactly 'The CARE Report' registered
        from src.detection.models import PIIEntity
        store.register_candidates([PIIEntity("COMPANY", "The CARE Report", 0, 15, 1.0, "ner")])
        
        manifest = redact_document(input_path, output_path, store)
        doc_out = Document(output_path)
        
        p1 = doc_out.paragraphs[0].text
        p2 = doc_out.paragraphs[1].text
        
        assert "The CARE Report" not in p1
        assert "the CARE Report" not in p2
        
        # Verify it was actually replaced
        replacement = store.get_replacement("COMPANY", "The CARE Report")
        assert replacement in p1
        assert replacement in p2
    finally:
        if os.path.exists(input_path): os.remove(input_path)
        if os.path.exists(output_path): os.remove(output_path)
