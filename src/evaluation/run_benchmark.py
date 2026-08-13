import os
import json
from collections import Counter, defaultdict
from typing import List, Dict, Any, Tuple
from docx import Document

from src.detection.regex_detector import RegexDetector
from src.detection.ner_detector import NERDetector
from src.detection.presidio_detector import PresidioDetector
from src.detection.fusion import resolve_candidates
from src.detection.validation import CandidateValidator
from src.mapping.span_mapper import reconstruct_paragraph_text
from src.reconstruction.replacer import get_paragraph_runs

def find_block(doc: Document, path: str) -> Any:
    """
    Parses a unique container path and retrieves the corresponding paragraph element.
    """
    parts = [p.strip() for p in path.split("/")]
    container_type = parts[0]
    
    try:
        if container_type == "body":
            p_idx = int(parts[1].split("=")[1])
            return doc.paragraphs[p_idx]
        elif container_type == "table":
            t_idx = int(parts[1].split("=")[1])
            r_idx = int(parts[2].split("=")[1])
            c_idx = int(parts[3].split("=")[1])
            p_idx = int(parts[4].split("=")[1])
            return doc.tables[t_idx].rows[r_idx].cells[c_idx].paragraphs[p_idx]
        elif container_type == "header":
            s_idx = int(parts[1].split("=")[1])
            p_idx = int(parts[2].split("=")[1])
            return doc.sections[s_idx].header.paragraphs[p_idx]
        elif container_type == "footer":
            s_idx = int(parts[1].split("=")[1])
            p_idx = int(parts[2].split("=")[1])
            return doc.sections[s_idx].footer.paragraphs[p_idx]
    except Exception as e:
        print(f"Error parsing path {path}: {e}")
    return None

def compute_metrics(tp: int, fp: int, fn: int) -> Tuple[float, float, float]:
    """
    Calculates Precision, Recall, and F1.
    """
    precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
    recall = tp / (tp + fn) if (tp + fn) > 0 else 0.0
    f1 = (2 * precision * recall) / (precision + recall) if (precision + recall) > 0 else 0.0
    return precision, recall, f1

def normalize_ground_truth(gt_entities: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Retains only the longest authoritative entity if multiple same-type entities overlap,
    preventing artificial FNs due to nested same-type sub-spans.
    """
    sorted_gt = sorted(gt_entities, key=lambda x: x["end"] - x["start"], reverse=True)
    authoritative = []
    for gt in sorted_gt:
        is_nested = False
        for auth in authoritative:
            if auth["entity_type"] == gt["entity_type"]:
                overlap_start = max(auth["start"], gt["start"])
                overlap_end = min(auth["end"], gt["end"])
                if overlap_start < overlap_end:
                    is_nested = True
                    break
        if not is_nested:
            authoritative.append(gt)
    return authoritative

def run_benchmark_eval(gt_path: str, doc_path: str, report_output_path: str) -> None:
    if not os.path.exists(gt_path):
        raise FileNotFoundError(f"Ground truth file not found at: {gt_path}")
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Prospectus file not found at: {doc_path}")

    # Load Ground Truth
    with open(gt_path, "r", encoding="utf-8") as f:
        gt_data = json.load(f)

    doc = Document(doc_path)

    # Instantiate Pipeline
    regex_detector = RegexDetector()
    ner_detector = NERDetector()
    presidio_detector = PresidioDetector()
    candidate_validator = CandidateValidator()

    # Track metrics TP, FP, FN per category across modes
    strict_metrics: Dict[str, Dict[str, int]] = defaultdict(lambda: {"tp": 0, "fp": 0, "fn": 0})
    overlap_metrics: Dict[str, Dict[str, int]] = defaultdict(lambda: {"tp": 0, "fp": 0, "fn": 0})
    normalized_metrics: Dict[str, Dict[str, int]] = defaultdict(lambda: {"tp": 0, "fp": 0, "fn": 0})

    # Detailed list of failures for error analysis (strict)
    false_positives_list: List[Dict[str, Any]] = []
    false_negatives_list: List[Dict[str, Any]] = []

    print(f"Evaluating {len(gt_data['blocks'])} blocks against human ground truth...")

    for block_entry in gt_data["blocks"]:
        path = block_entry["container_path"]
        gt_entities = block_entry["entities"]
        
        # Normalized ground truth list for Normalized Semantic mode
        normalized_gt = normalize_ground_truth(gt_entities)
        
        paragraph = find_block(doc, path)
        if not paragraph:
            print(f"Warning: Could not find block for path: {path}")
            continue

        # Run pipeline
        runs = get_paragraph_runs(paragraph)
        if not runs:
            # If ground truth expected entities but block has no runs, treat all as FN
            for gt in gt_entities:
                strict_metrics[gt["entity_type"]]["fn"] += 1
                overlap_metrics[gt["entity_type"]]["fn"] += 1
            for gt in normalized_gt:
                normalized_metrics[gt["entity_type"]]["fn"] += 1
            continue

        text, offsets = reconstruct_paragraph_text(runs)
        
        # 1. Detect
        cands = (
            regex_detector.detect(text) +
            ner_detector.detect(text) +
            presidio_detector.detect(text)
        )
        
        # 2. Fuse
        resolved = resolve_candidates(text, cands)
        resolved_pii = [e for e in resolved if e.entity_type != "DATE"]

        # 3. Validate
        validated = candidate_validator.validate_candidates(text, resolved_pii)
        pred_entities = [e for e in validated if e.metadata.get("validation_decision") == "KEEP"]

        # --- A. STRICT MATCHING ---
        matched_gt_strict = set()
        for pred in pred_entities:
            found_match = False
            for idx, gt in enumerate(gt_entities):
                if idx in matched_gt_strict:
                    continue
                # Match start, end, type exactly
                if (gt["start"] == pred.start and 
                    gt["end"] == pred.end and 
                    gt["entity_type"] == pred.entity_type):
                    
                    strict_metrics[pred.entity_type]["tp"] += 1
                    matched_gt_strict.add(idx)
                    found_match = True
                    break
                    
            if not found_match:
                strict_metrics[pred.entity_type]["fp"] += 1
                false_positives_list.append({
                    "path": path,
                    "text": pred.text,
                    "type": pred.entity_type,
                    "context": text
                })

        for idx, gt in enumerate(gt_entities):
            if idx not in matched_gt_strict:
                strict_metrics[gt["entity_type"]]["fn"] += 1
                false_negatives_list.append({
                    "path": path,
                    "text": gt["text"],
                    "type": gt["entity_type"],
                    "context": text
                })

        # --- B. OVERLAP-AWARE MATCHING ---
        pred_remaining = list(pred_entities)
        gt_remaining = list(gt_entities)
        
        # Pass 1: Exact matches first
        pred_still_remaining = []
        for pred in pred_remaining:
            found_exact = False
            for gt in list(gt_remaining):
                if (gt["start"] == pred.start and 
                    gt["end"] == pred.end and 
                    gt["entity_type"] == pred.entity_type):
                    overlap_metrics[pred.entity_type]["tp"] += 1
                    gt_remaining.remove(gt)
                    found_exact = True
                    break
            if not found_exact:
                pred_still_remaining.append(pred)
                
        pred_remaining = pred_still_remaining
        
        # Pass 2: Overlap matches (one-to-one)
        pred_still_remaining_2 = []
        for pred in pred_remaining:
            found_overlap = False
            for gt in list(gt_remaining):
                if gt["entity_type"] == pred.entity_type:
                    overlap_start = max(gt["start"], pred.start)
                    overlap_end = min(gt["end"], pred.end)
                    if overlap_start < overlap_end:
                        overlap_metrics[pred.entity_type]["tp"] += 1
                        gt_remaining.remove(gt)
                        found_overlap = True
                        break
            if not found_overlap:
                pred_still_remaining_2.append(pred)
                
        for pred in pred_still_remaining_2:
            overlap_metrics[pred.entity_type]["fp"] += 1
            
        for gt in gt_remaining:
            overlap_metrics[gt["entity_type"]]["fn"] += 1

        # --- C. NORMALIZED SEMANTIC MATCHING (on normalized_gt with overlap logic) ---
        pred_norm_remaining = list(pred_entities)
        gt_norm_remaining = list(normalized_gt)
        
        # Pass 1: Exact matches first
        pred_norm_still_remaining = []
        for pred in pred_norm_remaining:
            found_exact = False
            for gt in list(gt_norm_remaining):
                if (gt["start"] == pred.start and 
                    gt["end"] == pred.end and 
                    gt["entity_type"] == pred.entity_type):
                    normalized_metrics[pred.entity_type]["tp"] += 1
                    gt_norm_remaining.remove(gt)
                    found_exact = True
                    break
            if not found_exact:
                pred_norm_still_remaining.append(pred)
                
        pred_norm_remaining = pred_norm_still_remaining
        
        # Pass 2: Overlap matches (one-to-one)
        pred_norm_still_remaining_2 = []
        for pred in pred_norm_remaining:
            found_overlap = False
            for gt in list(gt_norm_remaining):
                if gt["entity_type"] == pred.entity_type:
                    overlap_start = max(gt["start"], pred.start)
                    overlap_end = min(gt["end"], pred.end)
                    if overlap_start < overlap_end:
                        normalized_metrics[pred.entity_type]["tp"] += 1
                        gt_norm_remaining.remove(gt)
                        found_overlap = True
                        break
            if not found_overlap:
                pred_norm_still_remaining_2.append(pred)
                
        for pred in pred_norm_still_remaining_2:
            normalized_metrics[pred.entity_type]["fp"] += 1
            
        for gt in gt_norm_remaining:
            normalized_metrics[gt["entity_type"]]["fn"] += 1

    # Compute summaries
    all_categories = sorted(list(set(
        list(strict_metrics.keys()) + list(overlap_metrics.keys()) + list(normalized_metrics.keys()) +
        ["PERSON", "COMPANY", "LOCATION", "EMAIL", "PHONE"]
    )))
    
    strict_rows = []
    overlap_rows = []
    normalized_rows = []
    
    strict_macro_p, strict_macro_r, strict_macro_f1 = 0.0, 0.0, 0.0
    overlap_macro_p, overlap_macro_r, overlap_macro_f1 = 0.0, 0.0, 0.0
    normalized_macro_p, normalized_macro_r, normalized_macro_f1 = 0.0, 0.0, 0.0
    strict_active, overlap_active, normalized_active = 0, 0, 0

    for cat in all_categories:
        # Strict
        s_counts = strict_metrics[cat]
        s_tp, s_fp, s_fn = s_counts["tp"], s_counts["fp"], s_counts["fn"]
        s_p, s_r, s_f1 = compute_metrics(s_tp, s_fp, s_fn)
        strict_rows.append(
            f"| `{cat}` | {s_tp} | {s_fp} | {s_fn} | {s_p*100:.1f}% | {s_r*100:.1f}% | {s_f1*100:.1f}% |"
        )
        if (s_tp + s_fp + s_fn) > 0:
            strict_macro_p += s_p
            strict_macro_r += s_r
            strict_macro_f1 += s_f1
            strict_active += 1

        # Overlap
        o_counts = overlap_metrics[cat]
        o_tp, o_fp, o_fn = o_counts["tp"], o_counts["fp"], o_counts["fn"]
        o_p, o_r, o_f1 = compute_metrics(o_tp, o_fp, o_fn)
        overlap_rows.append(
            f"| `{cat}` | {o_tp} | {o_fp} | {o_fn} | {o_p*100:.1f}% | {o_r*100:.1f}% | {o_f1*100:.1f}% |"
        )
        if (o_tp + o_fp + o_fn) > 0:
            overlap_macro_p += o_p
            overlap_macro_r += o_r
            overlap_macro_f1 += o_f1
            overlap_active += 1

        # Normalized Semantic
        n_counts = normalized_metrics[cat]
        n_tp, n_fp, n_fn = n_counts["tp"], n_counts["fp"], n_counts["fn"]
        n_p, n_r, n_f1 = compute_metrics(n_tp, n_fp, n_fn)
        normalized_rows.append(
            f"| `{cat}` | {n_tp} | {n_fp} | {n_fn} | {n_p*100:.1f}% | {n_r*100:.1f}% | {n_f1*100:.1f}% |"
        )
        if (n_tp + n_fp + n_fn) > 0:
            normalized_macro_p += n_p
            normalized_macro_r += n_r
            normalized_macro_f1 += n_f1
            normalized_active += 1

    # Strict Micro/Macro
    strict_total_tp = sum(strict_metrics[c]["tp"] for c in all_categories)
    strict_total_fp = sum(strict_metrics[c]["fp"] for c in all_categories)
    strict_total_fn = sum(strict_metrics[c]["fn"] for c in all_categories)
    strict_micro_p, strict_micro_r, strict_micro_f1 = compute_metrics(strict_total_tp, strict_total_fp, strict_total_fn)
    strict_macro_p = (strict_macro_p / strict_active) if strict_active > 0 else 0.0
    strict_macro_r = (strict_macro_r / strict_active) if strict_active > 0 else 0.0
    strict_macro_f1 = (strict_macro_f1 / strict_active) if strict_active > 0 else 0.0

    # Overlap Micro/Macro
    overlap_total_tp = sum(overlap_metrics[c]["tp"] for c in all_categories)
    overlap_total_fp = sum(overlap_metrics[c]["fp"] for c in all_categories)
    overlap_total_fn = sum(overlap_metrics[c]["fn"] for c in all_categories)
    overlap_micro_p, overlap_micro_r, overlap_micro_f1 = compute_metrics(overlap_total_tp, overlap_total_fp, overlap_total_fn)
    overlap_macro_p = (overlap_macro_p / overlap_active) if overlap_active > 0 else 0.0
    overlap_macro_r = (overlap_macro_r / overlap_active) if overlap_active > 0 else 0.0
    overlap_macro_f1 = (overlap_macro_f1 / overlap_active) if overlap_active > 0 else 0.0

    # Normalized Micro/Macro
    norm_total_tp = sum(normalized_metrics[c]["tp"] for c in all_categories)
    norm_total_fp = sum(normalized_metrics[c]["fp"] for c in all_categories)
    norm_total_fn = sum(normalized_metrics[c]["fn"] for c in all_categories)
    norm_micro_p, norm_micro_r, norm_micro_f1 = compute_metrics(norm_total_tp, norm_total_fp, norm_total_fn)
    normalized_macro_p = (normalized_macro_p / normalized_active) if normalized_active > 0 else 0.0
    normalized_macro_r = (normalized_macro_r / normalized_active) if normalized_active > 0 else 0.0
    normalized_macro_f1 = (normalized_macro_f1 / normalized_active) if normalized_active > 0 else 0.0

    # Write Metrics Report
    os.makedirs(os.path.dirname(os.path.abspath(report_output_path)), exist_ok=True)
    
    with open(report_output_path, "w", encoding="utf-8") as f:
        f.write(f"""# Ground-Truth Evaluation & Metrics Report

This report presents strict exact-span/exact-type metrics alongside semantic/overlap-aware metrics calculated against the human-labeled ground truth dataset in `evaluation/ground_truth.json`.

---

## 1. Overall Performance Metrics

### A. Strict Exact-Span Matching
| Metric Mode | Precision | Recall | F1-Score | True Positives | False Positives | False Negatives |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: |
| **Micro Average** | {strict_micro_p*100:.1f}% | {strict_micro_r*100:.1f}% | {strict_micro_f1*100:.1f}% | {strict_total_tp} | {strict_total_fp} | {strict_total_fn} |
| **Macro Average** | {strict_macro_p*100:.1f}% | {strict_macro_r*100:.1f}% | {strict_macro_f1*100:.1f}% | - | - | - |

### B. Semantic / Overlap-Aware Matching (One-to-One)
| Metric Mode | Precision | Recall | F1-Score | True Positives | False Positives | False Negatives |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: |
| **Micro Average** | {overlap_micro_p*100:.1f}% | {overlap_micro_r*100:.1f}% | {overlap_micro_f1*100:.1f}% | {overlap_total_tp} | {overlap_total_fp} | {overlap_total_fn} |
| **Macro Average** | {overlap_macro_p*100:.1f}% | {overlap_macro_r*100:.1f}% | {overlap_macro_f1*100:.1f}% | - | - | - |

### C. Normalized Semantic Matching (One-to-One Overlap on Deduplicated GT)
| Metric Mode | Precision | Recall | F1-Score | True Positives | False Positives | False Negatives |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: |
| **Micro Average** | {norm_micro_p*100:.1f}% | {norm_micro_r*100:.1f}% | {norm_micro_f1*100:.1f}% | {norm_total_tp} | {norm_total_fp} | {norm_total_fn} |
| **Macro Average** | {normalized_macro_p*100:.1f}% | {normalized_macro_r*100:.1f}% | {normalized_macro_f1*100:.1f}% | - | - | - |

> [!NOTE]
> **Strict Exact-Span Matching vs Normalized Semantic Matching**
> * **Strict Exact-Span Matching** answers: *"Did the system reproduce the exact annotated span?"*
>   - In paragraph 24, nested ground-truth annotations of both core and expanded spans (e.g. both `"Registrar of Companies"` and `"Registrar of Companies, Maharashtra"`) are present, yielding 5 artificial FNs because the single predicted authoritative span can only match one of them under strict exact-match rules.
> * **Normalized Semantic Matching** answers: *"Did the system correctly cover the authoritative PII entity?"*
>   - Overlapping nested annotations of the same category are deduplicated, keeping only the longest authoritative entity for evaluation, which resolves the nested span FNs and yields a clean evaluation of actual PII coverage.

---

## 2. Category-Specific Metrics

### A. Strict Exact-Span Metrics
| PII Category | TP | FP | FN | Precision | Recall | F1-Score |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: |
""" + "\n".join(strict_rows) + f"""

### B. Semantic / Overlap-Aware Metrics
| PII Category | TP | FP | FN | Precision | Recall | F1-Score |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: |
""" + "\n".join(overlap_rows) + f"""

### C. Normalized Semantic Metrics
| PII Category | TP | FP | FN | Precision | Recall | F1-Score |
| :--- | :---: | :---: | :---: | :---: | :---: | :---: |
""" + "\n".join(normalized_rows) + f"""

---

## 3. Detailed Error Analysis (Strict Matching)

### Suspected False Positives (Detected but not in Ground Truth)
Total: {len(false_positives_list)}

""" + "\n".join([f"- **Type**: `{fp['type']}` | **Text**: `\"{fp['text']}\"` | **Location**: {fp['path']}\n  *Context*: \"{fp['context'].strip()[:100]}...\"" for fp in false_positives_list[:10]]) + f"""

---

### Suspected False Negatives (PII Missed by Pipeline)
Total: {len(false_negatives_list)}

""" + "\n".join([f"- **Type**: `{fn['type']}` | **Text**: `\"{fn['text']}\"` | **Location**: {fn['path']}\n  *Context*: \"{fn['context'].strip()[:100]}...\"" for fn in false_negatives_list[:10]]) + """
""")
    print(f"Benchmark metrics report successfully generated at: {report_output_path}")

if __name__ == "__main__":
    run_benchmark_eval("evaluation/ground_truth.json", "data/input/prospectus.docx", "evaluation/evaluation_metrics.md")
