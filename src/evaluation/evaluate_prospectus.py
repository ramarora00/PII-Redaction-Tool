import os
import time
import re
from collections import Counter
from typing import List, Dict, Any, Tuple
from docx import Document

from src.detection.regex_detector import RegexDetector
from src.detection.ner_detector import NERDetector
from src.detection.presidio_detector import PresidioDetector
from src.detection.fusion import resolve_candidates
from src.detection.validation import CandidateValidator
from src.mapping.span_mapper import reconstruct_paragraph_text
from src.reconstruction.replacer import get_paragraph_runs

def run_evaluation(input_path: str, output_report_path: str) -> None:
    """
    Runs the complete evaluation pipeline over the real prospectus doc.
    Collects performance, detector, fusion, validation, type, and location metrics,
    saving a structured report to evaluation/prospectus_evaluation.md.
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Prospectus input file not found at: {input_path}")

    print(f"Starting evaluation of prospectus: {input_path}")
    
    # 1. Measure Document Loading Time
    start_load = time.perf_counter()
    doc = Document(input_path)
    end_load = time.perf_counter()
    load_time = end_load - start_load
    print(f"Document loaded in {load_time:.4f} seconds.")

    # Instantiate detectors
    regex_detector = RegexDetector()
    ner_detector = NERDetector()
    presidio_detector = PresidioDetector()
    candidate_validator = CandidateValidator()

    # Performance tracking
    total_regex_time = 0.0
    total_ner_time = 0.0
    total_presidio_time = 0.0
    total_fusion_time = 0.0
    total_validation_time = 0.0

    # Counts tracking
    scanned_counts = Counter()  # body_p, table_p, header_p, footer_p
    detector_raw_counts = Counter()  # regex, ner, presidio
    pii_type_counts = Counter()  # PERSON, EMAIL, etc.
    location_pii_counts = Counter()  # body, table, header, footer

    # Fusion metrics
    total_raw_candidates = 0
    total_resolved_candidates = 0

    # Validation metrics
    total_validated_candidates = 0
    total_rejected_candidates = 0
    validation_reasons = Counter()

    # Locations of PII tracking
    pii_entries: List[Dict[str, Any]] = []

    # Traverse all elements
    text_blocks: List[Tuple[Any, str, str]] = []

    # A. Body paragraphs
    for idx, p in enumerate(doc.paragraphs):
        text_blocks.append((p, "body", f"body / paragraph={idx}"))

    # B. Table cells
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    text_blocks.append((p, "table", f"table / table={t_idx} / row={r_idx} / cell={c_idx} / paragraph={p_idx}"))

    # C. Headers and footers
    for s_idx, section in enumerate(doc.sections):
        if section.header is not None:
            for p_idx, p in enumerate(section.header.paragraphs):
                text_blocks.append((p, "header", f"header / section={s_idx} / paragraph={p_idx}"))
        if section.footer is not None:
            for p_idx, p in enumerate(section.footer.paragraphs):
                text_blocks.append((p, "footer", f"footer / section={s_idx} / paragraph={p_idx}"))

    total_blocks = len(text_blocks)
    print(f"Scanned structure built. Total blocks to evaluate: {total_blocks}")

    start_eval_loop = time.perf_counter()

    for idx, (p, label, path) in enumerate(text_blocks):
        scanned_counts[label] += 1
        
        runs = get_paragraph_runs(p)
        if not runs:
            continue
            
        text, offsets = reconstruct_paragraph_text(runs)
        if not text.strip():
            continue

        # Progress indicator
        if idx > 0 and idx % 500 == 0:
            print(f"Progress: evaluated {idx}/{total_blocks} blocks...")

        # 1. Regex Detector
        t_start = time.perf_counter()
        regex_cands = regex_detector.detect(text)
        total_regex_time += (time.perf_counter() - t_start)
        detector_raw_counts["regex"] += len(regex_cands)

        # 2. spaCy NER Detector
        t_start = time.perf_counter()
        ner_cands = ner_detector.detect(text)
        total_ner_time += (time.perf_counter() - t_start)
        detector_raw_counts["ner"] += len(ner_cands)

        # 3. Presidio Detector
        t_start = time.perf_counter()
        presidio_cands = presidio_detector.detect(text)
        total_presidio_time += (time.perf_counter() - t_start)
        detector_raw_counts["presidio"] += len(presidio_cands)

        # 4. Fusion Layer
        all_cands = regex_cands + ner_cands + presidio_cands
        total_raw_candidates += len(all_cands)

        t_start = time.perf_counter()
        resolved = resolve_candidates(text, all_cands)
        total_fusion_time += (time.perf_counter() - t_start)
        total_resolved_candidates += len(resolved)

        # Filter out generic DATE (non-PII)
        resolved_pii = [e for e in resolved if e.entity_type != "DATE"]

        # 5. Validation Layer
        t_start = time.perf_counter()
        validated = candidate_validator.validate_candidates(text, resolved_pii)
        total_validation_time += (time.perf_counter() - t_start)

        for entity in validated:
            decision = entity.metadata.get("validation_decision", "KEEP")
            reason = entity.metadata.get("validation_reason", "GENUINE_PII")
            
            if decision == "KEEP":
                total_validated_candidates += 1
                pii_type_counts[entity.entity_type] += 1
                location_pii_counts[label] += 1

                # Extract sliding window context
                start_ctx = max(0, entity.start - 30)
                end_ctx = min(len(text), entity.end + 30)
                context = text[start_ctx:end_ctx].replace("\n", " ")
                if start_ctx > 0:
                    context = "..." + context
                if end_ctx < len(text):
                    context = context + "..."

                # Find matching original detector sources
                src_detectors = []
                for c in all_cands:
                    if c.start >= entity.start and c.end <= entity.end:
                        src_detectors.append(c.source)
                src_detectors = list(set(src_detectors))
                
                pii_entries.append({
                    "path": path,
                    "text": entity.text,
                    "type": entity.entity_type,
                    "context": context,
                    "detectors": " + ".join(src_detectors) if src_detectors else "unknown",
                    "reason": reason
                })
            else:
                total_rejected_candidates += 1
                validation_reasons[reason] += 1

    total_eval_time = time.perf_counter() - start_eval_loop
    print("Evaluation loop complete.")

    scanned_total = sum(scanned_counts.values())

    # Build Representative Examples
    rep_examples = []
    type_examples_count = Counter()
    
    for entry in pii_entries:
        etype = entry["type"]
        if type_examples_count[etype] < 3:
            type_examples_count[etype] += 1
            assessment = "SUSPECTED_TRUE_POSITIVE"
            
            rep_examples.append(f"""Type: {etype}
Text: {entry['text']}
Location: {entry['path']}
Detector: {entry['detectors']}
Final: {etype} (Reason: {entry['reason']})
Context: "{entry['context']}"
Assessment: {assessment}""")

    # Write Markdown Report
    os.makedirs(os.path.dirname(os.path.abspath(output_report_path)), exist_ok=True)
    
    with open(output_report_path, "w", encoding="utf-8") as f:
        f.write(f"""# Real Prospectus PII Evaluation Report

This report presents quantitative metrics and performance characteristics collected by running the PII redaction pipeline over `data/input/prospectus.docx`.

---

## 1. Execution Profile & Runtime Ratios

| Profiling Category | Duration (seconds) | Percentage of Total Execution |
| :--- | :---: | :---: |
| **Total Evaluation Execution** | {total_eval_time + load_time:.4f}s | 100.0% |
| Document Loading Time | {load_time:.4f}s | {(load_time / (total_eval_time + load_time)) * 100:.1f}% |
| Regex Detector Time | {total_regex_time:.4f}s | {(total_regex_time / (total_eval_time + load_time)) * 100:.1f}% |
| spaCy NER Detector Time | {total_ner_time:.4f}s | {(total_ner_time / (total_eval_time + load_time)) * 100:.1f}% |
| Presidio Analyzer Time | {total_presidio_time:.4f}s | {(total_presidio_time / (total_eval_time + load_time)) * 100:.1f}% |
| Fusion Layer Time | {total_fusion_time:.4f}s | {(total_fusion_time / (total_eval_time + load_time)) * 100:.1f}% |
| Validation Layer Time | {total_validation_time:.4f}s | {(total_validation_time / (total_eval_time + load_time)) * 100:.1f}% |

---

## 2. Document Traversal Summary

| Container Block Type | Scanned Blocks count | Percentage of Document |
| :--- | :---: | :---: |
| **Total Scanned Blocks** | {scanned_total} | 100.0% |
| Body Paragraphs | {scanned_counts['body']} | {(scanned_counts['body'] / scanned_total) * 100:.1f}% |
| Table Cell Paragraphs | {scanned_counts['table']} | {(scanned_counts['table'] / scanned_total) * 100:.1f}% |
| Section Headers | {scanned_counts['header']} | {(scanned_counts['header'] / scanned_total) * 100:.1f}% |
| Section Footers | {scanned_counts['footer']} | {(scanned_counts['footer'] / scanned_total) * 100:.1f}% |

---

## 3. Candidate Fusion & Redaction Summary

| Stage Metric | Candidate Count | Summary Description |
| :--- | :---: | :--- |
| **Raw Regex Candidates** | {detector_raw_counts['regex']} | Candidate matches from structured regex rules. |
| **Raw spaCy NER Candidates** | {detector_raw_counts['ner']} | Candidates from spaCy statistical model. |
| **Raw Presidio Candidates** | {detector_raw_counts['presidio']} | Candidates from Microsoft Presidio engine. |
| **Total Raw Candidates (Inputs)** | {total_raw_candidates} | Sum of all raw hits prior to fusion. |
| **Merged Duplicates & Conflicts** | {total_raw_candidates - total_resolved_candidates} | Candidates merged or dropped during overlapping deconfliction. |
| **Final Resolved Candidates** | {total_resolved_candidates} | Normalized candidates (including generic DATE). |
| **Candidates Rejected by Validator** | {total_rejected_candidates} | Suspicious false positives filtered by context validator. |
| **Final Redacted PII Spans** | {total_validated_candidates} | Verified PII entities (KEEP candidates). |

---

## 3.1 Candidate Validation Reason Breakdown

| Validation Reason | Suppression Count | Description |
| :--- | :---: | :--- |
| `GENERIC_DOCUMENT_TERM` | {validation_reasons['GENERIC_DOCUMENT_TERM']} | Common nouns / document metadata matching NER. |
| `ADDRESS_CONTEXT` | {validation_reasons['ADDRESS_CONTEXT']} | Address divisions misclassified as PERSON. |
| `TECHNICAL_STANDARD` | {validation_reasons['TECHNICAL_STANDARD']} | Technical standard numbers. |
| `FINANCIAL_CONTEXT` | {validation_reasons['FINANCIAL_CONTEXT']} | Numerical currency and scale units. |

---

## 4. Redacted PII Category Distribution

| PII Category | Final Redactions | Percentage of Redactions |
| :--- | :---: | :---: |
| **Total Redacted Spans** | {total_validated_candidates} | 100.0% |
| `PERSON` | {pii_type_counts['PERSON']} | {(pii_type_counts['PERSON'] / max(1, total_validated_candidates)) * 100:.1f}% |
| `COMPANY` | {pii_type_counts['COMPANY']} | {(pii_type_counts['COMPANY'] / max(1, total_validated_candidates)) * 100:.1f}% |
| `LOCATION` | {pii_type_counts['LOCATION']} | {(pii_type_counts['LOCATION'] / max(1, total_validated_candidates)) * 100:.1f}% |
| `EMAIL` | {pii_type_counts['EMAIL']} | {(pii_type_counts['EMAIL'] / max(1, total_validated_candidates)) * 100:.1f}% |
| `PHONE` | {pii_type_counts['PHONE']} | {(pii_type_counts['PHONE'] / max(1, total_validated_candidates)) * 100:.1f}% |
| `ADDRESS` | {pii_type_counts['ADDRESS']} | {(pii_type_counts['ADDRESS'] / max(1, total_validated_candidates)) * 100:.1f}% |
| `SSN` | {pii_type_counts['SSN']} | {(pii_type_counts['SSN'] / max(1, total_validated_candidates)) * 100:.1f}% |
| `CREDIT_CARD` | {pii_type_counts['CREDIT_CARD']} | {(pii_type_counts['CREDIT_CARD'] / max(1, total_validated_candidates)) * 100:.1f}% |
| `IP_ADDRESS` | {pii_type_counts['IP_ADDRESS']} | {(pii_type_counts['IP_ADDRESS'] / max(1, total_validated_candidates)) * 100:.1f}% |
| `DATE_OF_BIRTH` | {pii_type_counts['DATE_OF_BIRTH']} | {(pii_type_counts['DATE_OF_BIRTH'] / max(1, total_validated_candidates)) * 100:.1f}% |

---

## 5. Location Distribution of Redacted Spans

| Location Path Type | Redactions Count | Percentage of Redactions |
| :--- | :---: | :---: |
| **Total Redacted Spans** | {total_validated_candidates} | 100.0% |
| Body Paragraphs | {location_pii_counts['body']} | {(location_pii_counts['body'] / max(1, total_validated_candidates)) * 100:.1f}% |
| Table Cell Paragraphs | {location_pii_counts['table']} | {(location_pii_counts['table'] / max(1, total_validated_candidates)) * 100:.1f}% |
| Section Headers | {location_pii_counts['header']} | {(location_pii_counts['header'] / max(1, total_validated_candidates)) * 100:.1f}% |
| Section Footers | {location_pii_counts['footer']} | {(location_pii_counts['footer'] / max(1, total_validated_candidates)) * 100:.1f}% |

---

## 6. Representative Diagnostic Examples

```text
""" + "\n\n---\n\n".join(rep_examples) + """
```

---

## 7. Suspected Findings & Recommendations

### Suspected False Positives
- Numbers representing financial values or application indices incorrectly detected as `PHONE` or `CREDIT_CARD` (e.g. order numbers, folio references).
- General business terms (e.g. "Director", "Chairman", "Promoter") or common nouns incorrectly cataloged by spaCy/Presidio as `PERSON` or `COMPANY`.

### Suspected False Negatives
- Fragmented names across runs that fail dictionary NER checks.
- Address spans missing clean formatting boundaries.

### Architectural Recommendations for Milestone 8 (Tuning)
1. **Financial number suppressions**: Strengthen context rules to exclude sequences following financial abbreviations (e.g. `Rs.`, `Crore`, `Lakh`).
2. **False positive filters**: Implement custom lists to suppress generic nouns matched by NER engines.
""")
    print(f"Evaluation report generated successfully at: {output_report_path}")
